import base64
import hashlib
import hmac
import io
import json
import os
import random
import shutil
import string
import subprocess
import tempfile
import urllib.parse
import urllib.request
import uuid
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Annotated, Any

import fitz
import google.auth
import jwt
import pytesseract
from docx import Document
from fastapi import Depends, FastAPI, File, Form, Header, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from google.auth.transport.requests import Request as GoogleAuthRequest
from google.cloud import firestore as gcf
from google.cloud import storage as gcs
from pdf2docx import Converter
from PIL import Image


app = FastAPI(title="Sidekick Suites PDF Service")

allowed_origins = [
    origin.strip()
    for origin in os.getenv("ALLOWED_ORIGINS", "*").split(",")
    if origin.strip()
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

FONT_PATH = os.getenv("PDF_FONT_PATH", "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf")

# Standard PDF built-in font variants: base → (normal, bold, italic, bold-italic)
_FONT_VARIANTS: dict[str, dict[str, str]] = {
    "helv": {"": "helv", "b": "hebo", "i": "heit", "bi": "hebi"},
    "tiro": {"": "tiro", "b": "tibo", "i": "tiit", "bi": "tibi"},
    "cour": {"": "cour", "b": "cobo", "i": "coit", "bi": "cobi"},
}
_ALIGN_MAP = {"left": 0, "center": 1, "right": 2, "justify": 3}


def _resolve_font(family: str | None, bold: bool, italic: bool) -> tuple[str, str | None]:
    """Return (fontname, fontfile) for PyMuPDF insert_textbox."""
    fam = (family or "helv").lower()
    key = ("b" if bold else "") + ("i" if italic else "")
    if fam == "noto":
        ffile = FONT_PATH if Path(FONT_PATH).exists() else None
        return ("sidekick" if ffile else "helv"), ffile
    variants = _FONT_VARIANTS.get(fam, _FONT_VARIANTS["helv"])
    return variants.get(key, variants[""]), None


def _needs_unicode_font(text: str) -> bool:
    """Return True if text has characters outside WinAnsi (>U+00FF) like Vietnamese."""
    return any(ord(c) > 0xFF for c in text)


def _resolve_font_for_text(family: str | None, bold: bool, italic: bool, text: str) -> tuple[str, str | None]:
    """Like _resolve_font but auto-upgrades to NotoSans when text needs Unicode support."""
    font_name, font_file = _resolve_font(family, bold, italic)
    if _needs_unicode_font(text) and font_name != "sidekick":
        noto = Path(FONT_PATH)
        if noto.exists():
            return "sidekick", str(noto)
    return font_name, font_file
TESSERACT_LANG = os.getenv("TESSERACT_LANG", "eng+vie")
BUCKET_NAME = os.getenv("GCS_BUCKET", "sidekick-pdf-uploads")

# ── Auth config ──────────────────────────────────────────
ALLOWED_EMAIL_DOMAIN = os.getenv("ALLOWED_EMAIL_DOMAIN", "sidekick.vn")
JWT_SECRET = os.getenv("JWT_SECRET", "")
JWT_EXPIRE_DAYS = 30

GRAPH_TENANT_ID = os.getenv("GRAPH_TENANT_ID", "")
GRAPH_CLIENT_ID = os.getenv("GRAPH_CLIENT_ID", "")
GRAPH_CLIENT_SECRET = os.getenv("GRAPH_CLIENT_SECRET", "")
MAIL_FROM_ADDRESS = os.getenv("MAIL_FROM_ADDRESS", "noreply@sidekick.vn")

OTP_EXPIRE_MINUTES = 10
OTP_RESEND_COOLDOWN = 60
OTP_MAX_ATTEMPTS = 5

_firestore_client: gcf.Client | None = None


def _get_db() -> gcf.Client:
    global _firestore_client
    if _firestore_client is None:
        _firestore_client = gcf.Client(database=os.getenv("FIRESTORE_DB", "sidekick-suites"))
    return _firestore_client


def _generate_otp() -> str:
    return "".join(random.choices(string.digits, k=6))


def _hash_otp(code: str) -> str:
    return hashlib.sha256(code.encode()).hexdigest()


def _send_otp_email(to_email: str, code: str) -> None:
    # Get access token via client credentials flow
    token_data = urllib.parse.urlencode({
        "grant_type": "client_credentials",
        "client_id": GRAPH_CLIENT_ID,
        "client_secret": GRAPH_CLIENT_SECRET,
        "scope": "https://graph.microsoft.com/.default",
    }).encode()
    token_req = urllib.request.Request(
        f"https://login.microsoftonline.com/{GRAPH_TENANT_ID}/oauth2/v2.0/token",
        data=token_data, method="POST",
    )
    with urllib.request.urlopen(token_req, timeout=15) as resp:
        access_token = json.loads(resp.read())["access_token"]

    # Send via Graph API
    body = (
        f"Xin chào,\n\n"
        f"Mã đăng nhập Sidekick Suites của bạn là:\n\n"
        f"    {code}\n\n"
        f"Mã có hiệu lực trong {OTP_EXPIRE_MINUTES} phút.\n"
        f"Nếu bạn không yêu cầu mã này, hãy bỏ qua email này.\n\n"
        f"— Sidekick Suites"
    )
    mail_payload = json.dumps({
        "message": {
            "subject": f"[Sidekick Suites] Mã đăng nhập: {code}",
            "body": {"contentType": "Text", "content": body},
            "toRecipients": [{"emailAddress": {"address": to_email}}],
        }
    }).encode()
    mail_req = urllib.request.Request(
        f"https://graph.microsoft.com/v1.0/users/{MAIL_FROM_ADDRESS}/sendMail",
        data=mail_payload,
        headers={"Authorization": f"Bearer {access_token}", "Content-Type": "application/json"},
        method="POST",
    )
    with urllib.request.urlopen(mail_req, timeout=15) as resp:
        if resp.status not in (200, 202):
            raise Exception(f"Graph API returned {resp.status}")


def _issue_token(email: str) -> str:
    if not JWT_SECRET:
        raise HTTPException(status_code=500, detail="JWT_SECRET not configured")
    now = datetime.now(timezone.utc)
    payload = {"sub": email, "iat": now, "exp": now + timedelta(days=JWT_EXPIRE_DAYS)}
    return jwt.encode(payload, JWT_SECRET, algorithm="HS256")


def _verify_jwt(authorization: Annotated[str | None, Header()] = None) -> dict:
    if not JWT_SECRET:
        raise HTTPException(status_code=500, detail="JWT_SECRET not configured")
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Authentication required")
    token = authorization.removeprefix("Bearer ")
    try:
        return jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Session expired, please login again")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")
GS_QUALITY_MAP = {
    "screen": "/screen",
    "ebook": "/ebook",
    "printer": "/printer",
    "prepress": "/prepress",
}
DPI_MAP = {
    "screen": 72,
    "ebook": 150,
    "printer": 300,
    "prepress": 300,
}


def _read_pdf(upload: UploadFile) -> bytes:
    data = upload.file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty PDF upload")
    if not upload.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    return data


def _parse_pages(pages: str | None, total: int) -> list[int]:
    if not pages or not pages.strip():
        return list(range(total))
    selected: set[int] = set()
    for chunk in pages.split(","):
        part = chunk.strip()
        if not part:
            continue
        if "-" in part:
            start_s, end_s = part.split("-", 1)
            start = max(1, int(start_s))
            end = min(total, int(end_s))
            selected.update(range(start - 1, end))
        else:
            index = int(part) - 1
            if 0 <= index < total:
                selected.add(index)
    return sorted(selected)


def _safe_name(filename: str, suffix: str) -> str:
    stem = Path(filename).stem or "document"
    # HTTP headers must be ASCII; strip non-ASCII to avoid UnicodeEncodeError
    clean = "".join(ch for ch in stem if ch.isascii() and (ch.isalnum() or ch in ("-", "_", " "))).strip()
    return f"{clean or 'document'}{suffix}"


def _gs_available() -> bool:
    return shutil.which("gs") is not None


def _create_upload_url(filename: str) -> dict[str, str]:
    if not filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    client = gcs.Client()
    bucket = client.bucket(BUCKET_NAME)
    object_name = f"uploads/{uuid.uuid4()}/{filename}"
    blob = bucket.blob(object_name)
    creds, _ = google.auth.default(scopes=["https://www.googleapis.com/auth/cloud-platform"])
    creds.refresh(GoogleAuthRequest())
    sa_email = getattr(creds, "service_account_email", None) or os.getenv("SIGNING_SERVICE_ACCOUNT")
    if not sa_email:
        raise HTTPException(status_code=500, detail="Cannot resolve signing service account")
    upload_url = blob.generate_signed_url(
        version="v4",
        method="PUT",
        expiration=timedelta(minutes=30),
        content_type="application/pdf",
        service_account_email=sa_email,
        access_token=creds.token,
    )
    return {"upload_url": upload_url, "object_name": object_name}


def _build_gs_cmd(input_path: Path, output_path: Path, quality: str, grayscale: bool) -> list[str]:
    dpi = DPI_MAP[quality]
    cmd = [
        "gs",
        "-sDEVICE=pdfwrite",
        "-dNOPAUSE",
        "-dBATCH",
        "-dSAFER",
        f"-dPDFSETTINGS={GS_QUALITY_MAP[quality]}",
        "-dCompatibilityLevel=1.4",
        "-dCompressFonts=true",
        "-dSubsetFonts=true",
        "-dAutoRotatePages=/None",
        "-dDownsampleColorImages=true",
        "-dDownsampleGrayImages=true",
        "-dDownsampleMonoImages=true",
        "-dColorImageDownsampleType=/Average",
        "-dGrayImageDownsampleType=/Average",
        "-dMonoImageDownsampleType=/Subsample",
        f"-dColorImageResolution={dpi}",
        f"-dGrayImageResolution={dpi}",
        f"-dMonoImageResolution={min(dpi * 2, 300)}",
        "-dAutoFilterColorImages=false",
        "-dAutoFilterGrayImages=false",
        "-dColorImageFilter=/DCTEncode",
        "-dGrayImageFilter=/DCTEncode",
    ]
    if grayscale:
        cmd += ["-sColorConversionStrategy=Gray", "-dProcessColorModel=/DeviceGray"]
    cmd += [f"-sOutputFile={output_path}", str(input_path)]
    return cmd


def _hex_to_rgb01(hex_color: str | None) -> tuple[float, float, float]:
    value = (hex_color or "#000000").strip().lstrip("#")
    if len(value) != 6:
        return (0, 0, 0)
    try:
        return tuple(int(value[i : i + 2], 16) / 255 for i in (0, 2, 4))  # type: ignore[return-value]
    except ValueError:
        return (0, 0, 0)


def _expand_rect(rect: fitz.Rect, left: float, top: float, right: float, bottom: float) -> fitz.Rect:
    return fitz.Rect(rect.x0 + left, rect.y0 + top, rect.x1 + right, rect.y1 + bottom)


def _subset_pdf(data: bytes, page_indexes: list[int]) -> bytes:
    source = fitz.open(stream=data, filetype="pdf")
    output = fitz.open()
    for page_index in page_indexes:
        output.insert_pdf(source, from_page=page_index, to_page=page_index)
    out = output.tobytes(garbage=4, deflate=True)
    output.close()
    source.close()
    return out


def _text_lines(page: fitz.Page) -> list[str]:
    words = page.get_text("words")
    if not words:
        text = page.get_text("text").strip()
        return [line.strip() for line in text.splitlines() if line.strip()]
    words.sort(key=lambda word: (round(word[1] / 4) * 4, word[0]))
    lines: list[list[Any]] = []
    for word in words:
        y = word[1]
        if not lines or abs(lines[-1][0][1] - y) > 4:
            lines.append([word])
        else:
            lines[-1].append(word)
    return [" ".join(str(word[4]) for word in line).strip() for line in lines if line]


def _docx_from_text(data: bytes, page_indexes: list[int]) -> bytes:
    pdf = fitz.open(stream=data, filetype="pdf")
    document = Document()
    for offset, page_index in enumerate(page_indexes):
        page = pdf[page_index]
        lines = _text_lines(page)
        if not lines:
            document.add_paragraph("")
        for line in lines:
            document.add_paragraph(line)
        if offset < len(page_indexes) - 1:
            document.add_page_break()
    pdf.close()
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def _docx_from_ocr(data: bytes, page_indexes: list[int]) -> bytes:
    pdf = fitz.open(stream=data, filetype="pdf")
    document = Document()
    matrix = fitz.Matrix(2, 2)
    for offset, page_index in enumerate(page_indexes):
        page = pdf[page_index]
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        image = Image.open(io.BytesIO(pix.tobytes("png")))
        text = pytesseract.image_to_string(image, lang=TESSERACT_LANG)
        for line in text.splitlines():
            if line.strip():
                document.add_paragraph(line.strip())
        if offset < len(page_indexes) - 1:
            document.add_page_break()
    pdf.close()
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def _docx_from_layout(data: bytes, page_indexes: list[int]) -> bytes:
    subset = _subset_pdf(data, page_indexes)
    with tempfile.TemporaryDirectory() as tmp:
        pdf_path = Path(tmp) / "input.pdf"
        docx_path = Path(tmp) / "output.docx"
        pdf_path.write_bytes(subset)
        converter = Converter(str(pdf_path))
        try:
            converter.convert(str(docx_path))
        finally:
            converter.close()
        return docx_path.read_bytes()


def _action_rect(page: fitz.Page, action: dict[str, Any]) -> fitz.Rect:
    """Resolve the bounding rect for an action.

    New Fabric-based actions carry absolute PDF coordinates:
      pdfX / pdfY (Y from top, PyMuPDF style) / pdfWidth / pdfHeight
    Legacy actions use rect.{x,y,width,height} with y-from-bottom, or xPct/yPct.
    """
    page_width = page.rect.width
    page_height = page.rect.height

    # New coordinate format from Fabric.js (Y from top)
    if "pdfX" in action:
        x = float(action["pdfX"])
        y = float(action["pdfY"])          # Y from top
        w = float(action["pdfWidth"])
        h = float(action["pdfHeight"])
        return fitz.Rect(x, y, x + w, y + h)

    # Legacy rect format (Y from bottom)
    rect = action.get("rect")
    if rect:
        x = float(rect.get("x", 0))
        y = float(rect.get("y", 0))        # Y from bottom
        width = float(rect.get("width", 0))
        height = float(rect.get("height", 0))
        return fitz.Rect(x, page_height - y - height, x + width, page_height - y)

    # Legacy percent format
    x_pct = float(action.get("xPct", 0)) / 100
    y_pct = float(action.get("yPct", 0)) / 100
    width_pct = float(action.get("widthPct", 25)) / 100
    size = float(action.get("size", 14))
    x = page_width * x_pct
    y = page_height * y_pct
    return fitz.Rect(x, y, min(page_width, x + page_width * width_pct), min(page_height, y + size * 1.6))


def _original_rect_from_action(page: fitz.Page, action: dict[str, Any], fallback: fitz.Rect) -> fitz.Rect:
    # New Fabric format: originalRect is [x0, y0, x1, y1] in PyMuPDF (Y from top)
    orig = action.get("originalRect")
    if orig and isinstance(orig, (list, tuple)) and len(orig) == 4:
        return fitz.Rect(float(orig[0]), float(orig[1]), float(orig[2]), float(orig[3]))

    # Legacy format: rect dict with y from bottom
    rect = action.get("rect") or {}
    if not rect:
        return fallback
    page_height = page.rect.height
    x = float(rect.get("x", fallback.x0))
    width = float(rect.get("originalWidth") or rect.get("width") or fallback.width)
    height = float(rect.get("originalHeight") or rect.get("height") or fallback.height)
    y = float(rect.get("y", page_height - fallback.y1))
    replacement_height = float(rect.get("height") or fallback.height)
    original_y = y + replacement_height - height
    return fitz.Rect(x, page_height - original_y - height, x + width, page_height - original_y)


def _apply_edit_actions(data: bytes, actions: list[dict[str, Any]]) -> bytes:
    pdf = fitz.open(stream=data, filetype="pdf")

    for action in actions:
        page_index = int(action.get("page", 1)) - 1
        if page_index < 0 or page_index >= pdf.page_count:
            continue
        page = pdf[page_index]
        tool = action.get("tool", "")
        color = _hex_to_rgb01(action.get("color"))
        opacity = float(action.get("opacity", 1.0))
        rect = _action_rect(page, action)

        bold = bool(action.get("bold", False))
        italic = bool(action.get("italic", False))
        font_family = action.get("fontFamily") or "helv"
        align = _ALIGN_MAP.get(str(action.get("textAlign", "left")), 0)

        if tool == "replace":
            # Whiteout original text region then draw new text
            original_rect = _original_rect_from_action(page, action, rect)
            # Erase: 1pt sides, 2pt top for diacritics; bottom flush to avoid bleeding into block below
            page.add_redact_annot(_expand_rect(original_rect, -1, -2, 1, 0), fill=(1, 1, 1))
            page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)
            text = str(action.get("text", ""))
            font_name, font_file = _resolve_font_for_text(font_family, bold, italic, text)
            font_size = max(6.0, float(action.get("size") or original_rect.height * 0.75))
            # Use insert_text (point-based, no rect-fitting constraint) so text always renders.
            # Detected text from frontend is always single-line (newlines replaced with spaces).
            # Baseline ≈ y0 + 0.8*fontsize covers ascender for both Helvetica and NotoSans.
            baseline_y = original_rect.y0 + font_size * 0.82
            page.insert_text(
                fitz.Point(original_rect.x0, baseline_y),
                text,
                fontsize=font_size, fontname=font_name, fontfile=font_file,
                color=color,
            )

        elif tool == "text":
            text = str(action.get("text", ""))
            font_name, font_file = _resolve_font_for_text(font_family, bold, italic, text)
            font_size = float(action.get("size") or 14)
            page.insert_textbox(
                _expand_rect(rect, 0, -font_size * 0.45, max(rect.width * 0.2, font_size * 2), max(rect.height * 0.65, font_size)),
                text,
                fontsize=font_size, fontname=font_name, fontfile=font_file,
                color=color, align=align,
            )

        elif tool == "whiteout":
            page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1), width=0)

        elif tool == "highlight":
            annot = page.add_highlight_annot(rect)
            annot.set_colors(stroke=color)
            annot.set_opacity(min(1.0, opacity * 0.5))
            annot.update()

        elif tool == "rect":
            stroke = color if action.get("stroke", True) else None
            fill_color = _hex_to_rgb01(action.get("fillColor")) if action.get("fillColor") else None
            page.draw_rect(rect, color=stroke, fill=fill_color, width=float(action.get("strokeWidth", 1.5)))

        elif tool == "image":
            img_b64 = action.get("imageDataUrl", "")
            if img_b64:
                if "," in img_b64:
                    img_b64 = img_b64.split(",", 1)[1]
                try:
                    img_bytes = base64.b64decode(img_b64)
                    page.insert_image(rect, stream=img_bytes, keep_proportion=True)
                except Exception:
                    pass

        elif tool == "draw":
            # Freehand drawing exported as a PNG overlay
            img_b64 = action.get("imageDataUrl", "")
            if img_b64:
                if "," in img_b64:
                    img_b64 = img_b64.split(",", 1)[1]
                try:
                    img_bytes = base64.b64decode(img_b64)
                    page.insert_image(rect, stream=img_bytes, keep_proportion=False)
                except Exception:
                    pass

    out = pdf.tobytes(garbage=4, deflate=True)
    pdf.close()
    return out


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok", "ghostscript": str(_gs_available()).lower()}


@app.get("/")
def root() -> dict[str, Any]:
    return {"service": "Sidekick Suites PDF API", "ghostscript": _gs_available()}


@app.post("/auth/send-code")
async def send_code(email: str = Form(...)) -> dict:
    email = email.strip().lower()
    if not email.endswith(f"@{ALLOWED_EMAIL_DOMAIN}"):
        raise HTTPException(status_code=403, detail=f"Chỉ chấp nhận email @{ALLOWED_EMAIL_DOMAIN}")
    if not GRAPH_TENANT_ID or not GRAPH_CLIENT_ID or not GRAPH_CLIENT_SECRET:
        raise HTTPException(status_code=500, detail="Email service not configured")

    db = _get_db()
    doc_ref = db.collection("otp_codes").document(email)
    doc = doc_ref.get()
    now = datetime.now(timezone.utc)

    if doc.exists:
        sent_at = doc.to_dict().get("sent_at")
        if sent_at and (now - sent_at).total_seconds() < OTP_RESEND_COOLDOWN:
            remaining = OTP_RESEND_COOLDOWN - int((now - sent_at).total_seconds())
            raise HTTPException(status_code=429, detail=f"Vui lòng đợi {remaining}s trước khi gửi lại")

    code = _generate_otp()
    doc_ref.set({
        "code_hash": _hash_otp(code),
        "sent_at": now,
        "expires_at": now + timedelta(minutes=OTP_EXPIRE_MINUTES),
        "attempts": 0,
    })

    try:
        _send_otp_email(email, code)
    except Exception as exc:
        doc_ref.delete()
        raise HTTPException(status_code=500, detail=f"Không gửi được email: {type(exc).__name__}") from exc

    return {"ok": True, "message": f"Mã đã gửi đến {email}"}


@app.post("/auth/verify-code")
async def verify_code(email: str = Form(...), code: str = Form(...)) -> dict:
    email = email.strip().lower()
    code = code.strip()

    if not email.endswith(f"@{ALLOWED_EMAIL_DOMAIN}"):
        raise HTTPException(status_code=403, detail="Email không hợp lệ")

    db = _get_db()
    doc_ref = db.collection("otp_codes").document(email)
    doc = doc_ref.get()

    if not doc.exists:
        raise HTTPException(status_code=400, detail="Mã không hợp lệ hoặc đã hết hạn")

    data = doc.to_dict()
    now = datetime.now(timezone.utc)

    expires_at = data.get("expires_at")
    if not expires_at or expires_at < now:
        doc_ref.delete()
        raise HTTPException(status_code=400, detail="Mã đã hết hạn, vui lòng yêu cầu mã mới")

    attempts = data.get("attempts", 0)
    if attempts >= OTP_MAX_ATTEMPTS:
        doc_ref.delete()
        raise HTTPException(status_code=400, detail="Quá nhiều lần thử sai, vui lòng yêu cầu mã mới")

    if not hmac.compare_digest(data.get("code_hash", ""), _hash_otp(code)):
        doc_ref.update({"attempts": attempts + 1})
        remaining = OTP_MAX_ATTEMPTS - attempts - 1
        raise HTTPException(status_code=400, detail=f"Mã không đúng, còn {remaining} lần thử")

    doc_ref.delete()
    token = _issue_token(email)
    return {"token": token, "email": email, "expires_in_days": JWT_EXPIRE_DAYS}


@app.get("/upload-url")
def get_upload_url(filename: str, _user: dict = Depends(_verify_jwt)) -> dict[str, str]:
    return _create_upload_url(filename)


@app.post("/upload-url")
def post_upload_url(filename: str = Form(...), _user: dict = Depends(_verify_jwt)) -> dict[str, str]:
    return _create_upload_url(filename)


@app.post("/upload-pdf")
async def upload_pdf(file: UploadFile = File(...), _user: dict = Depends(_verify_jwt)) -> dict[str, str]:
    filename = file.filename or "upload.pdf"
    if not filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    client = gcs.Client()
    bucket = client.bucket(BUCKET_NAME)
    object_name = f"uploads/{uuid.uuid4()}/{filename}"
    blob = bucket.blob(object_name)
    blob.upload_from_file(file.file, content_type="application/pdf")
    return {"object_name": object_name}


@app.post("/compress-gcs")
async def compress_from_gcs(
    object_name: str = Form(...),
    quality: str = Form("ebook"),
    grayscale: bool = Form(False),
    _user: dict = Depends(_verify_jwt),
) -> FileResponse:
    if quality not in GS_QUALITY_MAP:
        raise HTTPException(status_code=400, detail=f"quality must be one of: {list(GS_QUALITY_MAP)}")
    if not _gs_available():
        raise HTTPException(status_code=503, detail="Ghostscript is not installed")
    try:
        client = gcs.Client()
        bucket = client.bucket(BUCKET_NAME)
        blob = bucket.blob(object_name)
        stem = Path(object_name).stem
        out_filename = f"{stem}_compressed.pdf"
        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "input.pdf"
            output_path = Path(tmp) / "compressed.pdf"
            blob.download_to_filename(str(input_path))
            original_size = input_path.stat().st_size
            result = subprocess.run(
                _build_gs_cmd(input_path, output_path, quality, grayscale),
                capture_output=True,
                timeout=1500,
                check=False,
            )
            try:
                blob.delete()
            except Exception:
                pass
            if result.returncode != 0 or not output_path.exists():
                final_path = Path(tempfile.gettempdir()) / out_filename
                shutil.copy2(input_path, final_path)
                return FileResponse(
                    path=str(final_path),
                    media_type="application/pdf",
                    filename=out_filename,
                    headers={
                        "X-Original-Size": str(original_size),
                        "X-Compressed-Size": str(original_size),
                        "X-Size-Reduction": "0",
                        "X-Compress-Fallback": "ghostscript_failed_returned_original",
                    },
                )
            compressed_size = output_path.stat().st_size
            final_path = Path(tempfile.gettempdir()) / out_filename
            if compressed_size >= original_size:
                shutil.copy2(input_path, final_path)
                compressed_size = original_size
            else:
                shutil.copy2(output_path, final_path)
            reduction_pct = round((1 - compressed_size / original_size) * 100, 1) if original_size else 0
            return FileResponse(
                path=str(final_path),
                media_type="application/pdf",
                filename=out_filename,
                headers={
                    "X-Original-Size": str(original_size),
                    "X-Compressed-Size": str(compressed_size),
                    "X-Size-Reduction": str(reduction_pct),
                },
            )
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"compress-gcs internal error: {type(exc).__name__}: {str(exc)[:300]}",
        ) from exc


@app.post("/pdf/analyze-text")
async def analyze_text(
    file: UploadFile = File(...),
    pages: str | None = Form(None),
    _user: dict = Depends(_verify_jwt),
) -> JSONResponse:
    data = _read_pdf(file)
    pdf = fitz.open(stream=data, filetype="pdf")
    selected = _parse_pages(pages, pdf.page_count)
    result = []
    for page_index in selected:
        page = pdf[page_index]
        blocks = []
        for block in page.get_text("blocks"):
            x0, y0, x1, y1, text, *_ = block
            if str(text).strip():
                blocks.append({"rect": [x0, y0, x1, y1], "text": str(text).strip()})
        result.append({"page": page_index + 1, "width": page.rect.width, "height": page.rect.height, "blocks": blocks})
    pdf.close()
    return JSONResponse({"pages": result})


@app.post("/pdf/edit-text")
async def edit_text(
    file: UploadFile = File(...),
    actions: str = Form("[]"),
    _user: dict = Depends(_verify_jwt),
) -> StreamingResponse:
    data = _read_pdf(file)
    try:
        parsed_actions = json.loads(actions.lstrip("\ufeff"))
        if not isinstance(parsed_actions, list):
            raise ValueError("actions must be a list")
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=f"Invalid actions JSON: {exc}") from exc

    output = _apply_edit_actions(data, parsed_actions)
    filename = _safe_name(file.filename, "_edited.pdf")
    return StreamingResponse(
        io.BytesIO(output),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/pdf/convert-word")
async def convert_word(
    file: UploadFile = File(...),
    pages: str | None = Form(None),
    mode: str = Form("layout"),
    _user: dict = Depends(_verify_jwt),
) -> StreamingResponse:
    data = _read_pdf(file)
    pdf = fitz.open(stream=data, filetype="pdf")
    selected = _parse_pages(pages, pdf.page_count)
    pdf.close()
    if not selected:
        raise HTTPException(status_code=400, detail="No valid pages selected")

    try:
        if mode == "text":
            output = _docx_from_text(data, selected)
        elif mode == "ocr":
            output = _docx_from_ocr(data, selected)
        else:
            output = _docx_from_layout(data, selected)
    except Exception as exc:
        if mode == "layout":
            output = _docx_from_text(data, selected)
        else:
            raise HTTPException(status_code=500, detail=str(exc)) from exc

    filename = _safe_name(file.filename, ".docx")
    return StreamingResponse(
        io.BytesIO(output),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
